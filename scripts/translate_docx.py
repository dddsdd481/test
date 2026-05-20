#!/usr/bin/env python3
"""Translate C1.docx to Moroccan Darija (Arabic script) preserving layout."""
import os
import re
import json
import sys
import time
from copy import deepcopy

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI

INPUT = "/vercel/share/v0-project/C1.docx"
OUTPUT = "/vercel/share/v0-project/C1-ar.docx"

GATEWAY_KEY = os.environ["AI_GATEWAY_API_KEY"]
client = OpenAI(
    base_url="https://ai-gateway.vercel.sh/v1",
    api_key=GATEWAY_KEY,
)
MODEL = "anthropic/claude-haiku-4.5"

SYSTEM = """You are a professional translator. Translate text from English into Moroccan Darija written in ARABIC SCRIPT (not Latin/Arabizi).

Rules:
- Use natural spoken Moroccan Darija, NOT Modern Standard Arabic. Use words like: كاين، بزاف، دابا، هاد، ديال، فهاد، علاش، باش، خاصك، ماشي، كنشوفو، نديرو، نقدرو، عاد، حيت، ولا، ولكن، غير، تما، هنا، هاكدا.
- Keep these technical terms in ENGLISH (Latin script): Machine Learning, Supervised Learning, Unsupervised Learning, Regression, Classification, Linear Regression, Logistic Regression, Gradient Descent, Cost Function, Loss Function, Feature Scaling, Regularization, Overfitting, Underfitting, Vectorization, NumPy, Python, Jupyter, neural network, clustering, training set, test set, learning rate, bias, variance, sigmoid, decision boundary, hyperparameter, epoch, batch, mini-batch, lab, optional lab, week, model, dataset, feature, label, parameter, weight.
- Keep ALL formulas, math symbols, code, variable names, file names, lab titles, week numbers, and identifiers UNCHANGED (e.g. w, b, x, y, J(w,b), f(x), alpha, P22, P30).
- When introducing a technical term for the first time in a passage, you may add a SHORT Darija clarification in parentheses. Don't repeat clarifications.
- Preserve numbers, punctuation style, and inline math.
- Output ONLY the translation. No preface, no explanation, no quotes around it.
- If the input is a heading or short label, translate it concisely.
- If input is just a code/identifier/number/formula with no real prose, return it unchanged.

You will receive a JSON array of strings. Output a JSON array of the same length with translations in the same order. Output ONLY the JSON array, nothing else."""

def translate_batch(texts):
    if not texts:
        return []
    payload = json.dumps(texts, ensure_ascii=False)
    for attempt in range(4):
        try:
            r = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": SYSTEM},
                    {"role": "user", "content": payload},
                ],
                temperature=0.3,
                max_tokens=8000,
            )
            content = r.choices[0].message.content.strip()
            # strip code fences if any
            content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content, flags=re.MULTILINE).strip()
            # find first [ and last ]
            start = content.find("[")
            end = content.rfind("]")
            if start >= 0 and end > start:
                content = content[start : end + 1]
            arr = json.loads(content)
            if len(arr) != len(texts):
                raise ValueError(f"length mismatch: {len(arr)} vs {len(texts)}")
            return [str(x) for x in arr]
        except Exception as e:
            print(f"  retry {attempt+1}: {e}", file=sys.stderr)
            time.sleep(2 + attempt * 2)
    # fallback: translate one by one
    out = []
    for t in texts:
        try:
            sub = translate_batch([t]) if len(texts) > 1 else [t]
            out.append(sub[0])
        except Exception:
            out.append(t)
    return out


def should_skip(text):
    s = text.strip()
    if not s:
        return True
    # Pure number / page marker like "P22"
    if re.fullmatch(r"P\d+", s):
        return True
    # Pure formula/identifier/code
    if re.fullmatch(r"[\d\s\.\,\-\+\*/=<>\(\)\[\]\{\}\^_\|\\:%a-zA-Z]{1,8}", s) and not re.search(r"[a-zA-Z]{4,}", s):
        return True
    return False


def set_paragraph_rtl(p):
    pPr = p._p.get_or_add_pPr()
    # bidi
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    # right alignment for non-centered
    jc = pPr.find(qn("w:jc"))
    cur = jc.get(qn("w:val")) if jc is not None else None
    if cur not in ("center",):
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "right")


def set_run_rtl_arabic(run):
    rPr = run._r.get_or_add_rPr()
    # rtl
    rtl = rPr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)
    # set complex script font
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:cs"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:ascii"), "Arial")


def replace_paragraph_text(p, new_text):
    """Replace paragraph text, keeping the first run's formatting."""
    runs = p.runs
    if not runs:
        # add a run
        r = p.add_run(new_text)
        set_run_rtl_arabic(r)
        return
    # Keep first run, clear others
    first = runs[0]
    first.text = new_text
    set_run_rtl_arabic(first)
    for r in runs[1:]:
        r.text = ""
        # remove element to keep doc clean
        r._r.getparent().remove(r._r)


def collect_paragraphs(doc):
    items = []  # (paragraph_obj, text)
    for p in doc.paragraphs:
        if p.text.strip():
            items.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        items.append(p)
    return items


def main():
    doc = docx.Document(INPUT)
    paragraphs = collect_paragraphs(doc)
    print(f"Found {len(paragraphs)} non-empty paragraphs")

    # Build translation list (skip ones we shouldn't translate)
    to_translate_idx = []
    to_translate_text = []
    for i, p in enumerate(paragraphs):
        t = p.text
        if should_skip(t):
            continue
        to_translate_idx.append(i)
        to_translate_text.append(t)

    print(f"Will translate {len(to_translate_text)} paragraphs")

    # Batch by character count
    BATCH_CHAR_LIMIT = 4000
    translations = {}
    batch_idx = []
    batch_txt = []
    cur_chars = 0
    batch_num = 0

    def flush():
        nonlocal batch_num, cur_chars
        if not batch_txt:
            return
        batch_num += 1
        print(f"  batch {batch_num}: {len(batch_txt)} items, {cur_chars} chars", flush=True)
        result = translate_batch(batch_txt)
        for idx, tr in zip(batch_idx, result):
            translations[idx] = tr
        batch_idx.clear()
        batch_txt.clear()
        cur_chars = 0

    for idx, txt in zip(to_translate_idx, to_translate_text):
        if cur_chars + len(txt) > BATCH_CHAR_LIMIT and batch_txt:
            flush()
        batch_idx.append(idx)
        batch_txt.append(txt)
        cur_chars += len(txt)
    flush()

    # Apply translations
    for i, p in enumerate(paragraphs):
        if i in translations:
            replace_paragraph_text(p, translations[i])
        # set RTL on every paragraph that contains text (translated or kept)
        if p.text.strip():
            set_paragraph_rtl(p)
            for r in p.runs:
                set_run_rtl_arabic(r)

    # Set section to RTL
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = sectPr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            sectPr.append(bidi)

    doc.save(OUTPUT)
    print(f"Saved {OUTPUT}")


if __name__ == "__main__":
    main()
