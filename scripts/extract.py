#!/usr/bin/env python3
"""Extract all translatable paragraphs from C1.docx into JSON."""
import json
import re
import docx

INPUT = "/vercel/share/v0-project/C1.docx"
OUT = "/vercel/share/v0-project/scripts/extracted.json"


def should_skip(text):
    s = text.strip()
    if not s:
        return True
    if re.fullmatch(r"P\d+", s):
        return True
    # Pure short identifier/formula
    if len(s) <= 6 and not re.search(r"[a-zA-Z]{4,}", s):
        return True
    return False


def collect(doc):
    items = []
    for p_idx, p in enumerate(doc.paragraphs):
        if p.text.strip():
            items.append(("p", p_idx, None, None, None, p.text))
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for cp_idx, p in enumerate(cell.paragraphs):
                    if p.text.strip():
                        items.append(("t", t_idx, r_idx, c_idx, cp_idx, p.text))
    return items


def main():
    doc = docx.Document(INPUT)
    items = collect(doc)
    data = []
    for it in items:
        kind, a, b, c, d, text = it
        data.append(
            {
                "loc": [kind, a, b, c, d],
                "text": text,
                "skip": should_skip(text),
            }
        )
    with open(OUT, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=1)
    print(f"Wrote {len(data)} items to {OUT}")
    print(f"To translate: {sum(1 for d in data if not d['skip'])}")
    total_chars = sum(len(d['text']) for d in data if not d['skip'])
    print(f"Total chars to translate: {total_chars}")


if __name__ == "__main__":
    main()
